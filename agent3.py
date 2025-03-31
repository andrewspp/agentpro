#!/usr/bin/env python3
"""
Agent 3: Synthèse et Rapport PDF

Ce script génère un rapport structuré au format PDF à partir des analyses 
réalisées par les agents 1 et 2. Il assure une progression logique des idées,
génère les sections complémentaires et met en forme le rapport final.

Usage:
    python agent3.py agent1_output agent2_output user_prompt [--model modele] [--backend backend]

Arguments:
    agent1_output: Fichier JSON généré par l'agent 1
    agent2_output: Fichier JSON généré par l'agent 2
    user_prompt: Prompt utilisateur original
    --model: Modèle LLM à utiliser (défaut: gemma3:27b)
    --backend: Backend LLM ('ollama' ou 'gemini', défaut: 'ollama')
"""

import argparse
import json
import logging
import os
import sys
import base64
from datetime import datetime
import markdown
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML, CSS
import re
import pandas as pd
import io

# Importation conditionnelle pour gérer l'absence de modules
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import html2text
    DOCX_AVAILABLE = True
except ImportError as e:
    logging.warning(f"Module python-docx non disponible. La génération de document Word sera désactivée. Erreur: {e}")
    logging.warning("Pour activer cette fonctionnalité, exécutez: pip install python-docx html2text")

# Importation du module llm_utils
from llm_utils import call_llm

# Configuration du logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("agent3.log"),
        logging.StreamHandler(sys.stderr)
    ]
)
logger = logging.getLogger("agent3")

# ======================================================
# Fonctions utilitaires pour les templates
# ======================================================

def basename_filter(path):
    """
    Filtre Jinja2 pour obtenir le nom de base d'un chemin.
    
    Args:
        path: Chemin à traiter
        
    Returns:
        str: Nom de base du chemin
    """
    return os.path.basename(path) if path else path

def preprocess_markdown(text):
    """
    Prétraite le texte Markdown pour assurer une meilleure compatibilité avec les convertisseurs.
    
    Args:
        text: Texte Markdown à prétraiter
        
    Returns:
        str: Texte Markdown prétraité
    """
    if not text:
        return ""
    
    # Diviser le texte en lignes
    lines = text.split('\n')
    processed_lines = []
    
    in_list = False
    list_indent = 0
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Correction des titres - s'assurer qu'il y a un espace après le #
        if stripped and stripped[0] == '#':
            for level in range(6, 0, -1):  # Vérifier les niveaux de titre de h6 à h1
                tag = '#' * level
                if stripped.startswith(tag) and (len(stripped) == level or stripped[level] != ' '):
                    line = tag + ' ' + stripped[level:].lstrip()
                    break
                    
        # Pré-traitement spécial pour les titres générés par le LLM
        # Souvent, le LLM peut générer des titres comme "# 1. TITRE" ou "## Titre"
        # sans que ce soit correctement interprété comme Markdown
        if stripped.startswith('# ') or stripped.startswith('## ') or stripped.startswith('### '):
            # Assurons-nous que ces titres soient sans symboles #
            level = 0
            while stripped[level] == '#':
                level += 1
            line = stripped[level:].strip()  # Supprime les # et les espaces
        
        # Correction des listes à puces
        if stripped.startswith('-') and not stripped.startswith('- '):
            spaces_before = len(line) - len(line.lstrip())
            line = ' ' * spaces_before + '- ' + stripped[1:].lstrip()
            in_list = True
            list_indent = spaces_before
            
        # Correction des listes numériques
        elif re.match(r'^\d+\.\S', stripped):
            num, rest = re.match(r'^(\d+\.)(\S.*)', stripped).groups()
            spaces_before = len(line) - len(line.lstrip())
            line = ' ' * spaces_before + num + ' ' + rest
            in_list = True
            list_indent = spaces_before
            
        # Gestion des éléments de liste qui continuent sur plusieurs lignes
        elif in_list and stripped and lines[i-1].strip() and len(line) - len(line.lstrip()) >= list_indent:
            # C'est la continuation d'un élément de liste
            pass
        elif stripped:
            # Ligne non vide qui n'est pas une continuation de liste
            in_list = False
        
        # Correction des tableaux mal formatés
        if '|' in line:
            # Assurer que les cellules de tableau ont des espaces appropriés
            parts = line.split('|')
            line = '|'.join([p.strip() for p in parts])
        
        processed_lines.append(line)
    
    return '\n'.join(processed_lines)

def markdown_filter(text):
    """
    Filtre Jinja2 amélioré pour convertir du Markdown en HTML avec support avancé.
    
    Args:
        text: Texte Markdown à convertir
        
    Returns:
        str: HTML généré
    """
    if not text:
        return ""
    
    try:
        # Prétraitement du texte Markdown
        text = preprocess_markdown(text)
        
        # Nettoyer les caractères # qui ne sont pas des titres (au milieu du texte)
        # Mais conserver ceux qui sont au début des lignes pour les convertir en titres
        lines = text.split('\n')
        for i, line in enumerate(lines):
            # Si le caractère # n'est pas au début de la ligne après des espaces, l'échapper
            if '#' in line and not line.lstrip().startswith('#'):
                # Remplacer les # qui ne sont pas au début de la ligne par \#
                parts = []
                j = 0
                while j < len(line):
                    if line[j] == '#' and (j == 0 or line[j-1] != '\\'):
                        # Vérifier si ce # est au début d'une ligne ou après des espaces
                        prefix = line[:j].rstrip()
                        if prefix:  # S'il y a du texte avant, échapper le #
                            parts.append(line[:j] + '\\')
                            parts.append(line[j:])
                            break
                    j += 1
                if parts:
                    lines[i] = ''.join(parts)
        
        text = '\n'.join(lines)
        
        # Utilisation d'extensions supplémentaires pour améliorer le rendu
        html = markdown.markdown(text, extensions=[
            'fenced_code',       # Pour les blocs de code ```
            'codehilite',        # Coloration syntaxique des blocs de code
            'tables',            # Pour les tableaux
            'nl2br',             # Conversion des sauts de ligne
            'sane_lists',        # Listes mieux formatées
            'smarty',            # Guillemets intelligents et tirets
            'attr_list',         # Attributs pour les éléments
            'def_list',          # Listes de définition
            'footnotes',         # Notes de bas de page
            'md_in_html'         # Permet le Markdown dans les balises HTML
        ])
        
        # Post-traitement pour corriger certains problèmes courants
        # Assurer que les listes ont les bonnes classes CSS
        html = html.replace('<ul>', '<ul class="report-list">')
        html = html.replace('<ol>', '<ol class="report-list">')
        
        # Assurer que les paragraphes dans les éléments de liste sont bien formatés
        html = html.replace('<li><p>', '<li>')
        html = html.replace('</p></li>', '</li>')
        
        return html
    except Exception as e:
        logger.error(f"Erreur lors de la conversion Markdown: {e}")
        # Fallback simple en cas d'erreur
        return f"<p>{text}</p>"

def csv_to_html_table(csv_data, max_rows=10):
    """
    Convertit les données CSV en tableau HTML.
    
    Args:
        csv_data: Données au format CSV sous forme de chaîne
        max_rows: Nombre maximum de lignes à afficher
        
    Returns:
        str: Tableau HTML ou message d'erreur
    """
    if not csv_data or len(csv_data.strip()) == 0:
        return ""
    
    try:
        # Parser les données CSV
        df = pd.read_csv(io.StringIO(csv_data))
        
        # Limiter le nombre de lignes
        if len(df) > max_rows:
            df = df.head(max_rows)
            footer = f"<p><em>Affichage des {max_rows} premières lignes sur {len(df)} au total.</em></p>"
        else:
            footer = ""
            
        # Convertir en HTML
        table_html = df.to_html(index=False, classes="csv-data-table")
        
        return f"<div class='csv-data-container'>{table_html}{footer}</div>"
    
    except Exception as e:
        logger.error(f"Erreur lors de la conversion des données CSV en HTML: {e}")
        return f"<div class='error'>Erreur de conversion des données CSV: {str(e)}</div>"

def parse_csv_for_summary(csv_data):
    """
    Extrait des informations résumées des données CSV.
    
    Args:
        csv_data: Données au format CSV sous forme de chaîne
        
    Returns:
        dict: Résumé des données ou dictionnaire vide en cas d'erreur
    """
    if not csv_data or len(csv_data.strip()) == 0:
        return {}
    
    try:
        # Parser les données CSV
        df = pd.read_csv(io.StringIO(csv_data))
        
        # Créer un résumé
        summary = {
            "num_rows": len(df),
            "num_cols": len(df.columns),
            "columns": list(df.columns),
            "numeric_stats": {}
        }
        
        # Ajouter des statistiques pour les colonnes numériques
        for col in df.select_dtypes(include=['number']).columns:
            summary["numeric_stats"][col] = {
                "min": float(df[col].min()) if not pd.isna(df[col].min()) else None,
                "max": float(df[col].max()) if not pd.isna(df[col].max()) else None,
                "mean": float(df[col].mean()) if not pd.isna(df[col].mean()) else None,
                "median": float(df[col].median()) if not pd.isna(df[col].median()) else None,
            }
            
        return summary
    
    except Exception as e:
        logger.error(f"Erreur lors de l'analyse CSV pour résumé: {e}")
        return {}

# ======================================================
# Fonctions de traitement des images et visualisations
# ======================================================

def save_images(visualisations, img_dir):
    """
    Sauvegarde les visualisations à partir de base64 et retourne une liste de dictionnaires.
    Prend en charge à la fois les visualisations et les tables de régression.
    
    Args:
        visualisations: Liste des métadonnées des visualisations
        img_dir: Répertoire où sauvegarder les images
        
    Returns:
        Liste des informations sur les images sauvegardées
    """
    os.makedirs(img_dir, exist_ok=True)
    image_infos = []  # Pour stocker plus d'infos

    for i, vis in enumerate(visualisations):
        # Déterminer l'identifiant et le titre
        if 'filename' in vis:
            base_filename = vis.get("filename", f"figure_{i+1}.png")
            vis_id = os.path.splitext(base_filename)[0]
        else:
            vis_id = vis.get("id", f"item_{i+1}")
            base_filename = vis_id + '.png'
        
        # Déterminer le type de visualisation
        vis_type = "visualisation"
        if 'type' in vis and vis['type'] == 'regression_table':
            vis_type = "regression_table"
            
        # Déterminer le titre
        title = vis.get("title", vis_id.replace('_', ' ').capitalize())
        
        # Nom de fichier standardisé
        img_filename = f"{vis_id}.png"
        img_path = os.path.join(img_dir, img_filename)

        if "base64" in vis:
            try:
                img_data = base64.b64decode(vis["base64"])
                with open(img_path, "wb") as f:
                    f.write(img_data)
                
                # Journaliser plus d'informations pour le débogage
                file_size = os.path.getsize(img_path)
                logger.info(f"Sauvegarde de l'image {img_path} avec taille: {file_size} octets")
                
                # Extraire les données CSV si disponibles
                csv_data = vis.get("csv_data", "")
                csv_summary = parse_csv_for_summary(csv_data) if csv_data else {}
                
                # Extraire l'interprétation détaillée si disponible (pour les régressions)
                detailed_interpretation = vis.get("detailed_interpretation", "")
                
                # Créer l'objet d'information avec les métadonnées
                image_info = {
                    "filename": img_filename,
                    "path": img_path,
                    "title": title,
                    "type": vis_type,
                    "interpretation": vis.get("interpretation", ""),
                    "detailed_interpretation": detailed_interpretation,  # Ajouter l'interprétation détaillée
                    "metadata": vis.get("metadata", {}),
                    "data": vis.get("data", {}),
                    "csv_data": csv_data,
                    "csv_summary": csv_summary,
                    "csv_html": csv_to_html_table(csv_data) if csv_data else "",
                    "size": file_size
                }
                
                image_infos.append(image_info)
                logger.info(f"Image sauvegardée: {img_path} avec titre: {image_info['title']}")
                if csv_data:
                    logger.info(f"  -> Données CSV trouvées pour: {img_filename} ({len(csv_data)} caractères)")
                if detailed_interpretation:
                    logger.info(f"  -> Interprétation détaillée de {len(detailed_interpretation)} caractères transférée")
            except Exception as e:
                logger.error(f"Erreur lors de la sauvegarde de l'image {vis_id}: {e}")
                logger.error(f"Contenu de 'base64': {vis.get('base64', '')[:50]}...")
        else:
            logger.warning(f"Pas de données base64 pour l'image {vis_id}")
            # Si path est présent, essayer de copier l'image
            if 'path' in vis and os.path.exists(vis['path']):
                try:
                    import shutil
                    shutil.copy(vis['path'], img_path)
                    
                    # Extraire les données CSV si disponibles
                    csv_data = vis.get("csv_data", "")
                    csv_summary = parse_csv_for_summary(csv_data) if csv_data else {}
                    
                    # Extraire l'interprétation détaillée (pour les régressions)
                    detailed_interpretation = vis.get("detailed_interpretation", "")
                    
                    image_info = {
                        "filename": img_filename,
                        "path": img_path,
                        "title": title,
                        "type": vis_type,
                        "interpretation": vis.get("interpretation", ""),
                        "detailed_interpretation": detailed_interpretation,  # Ajouter l'interprétation détaillée
                        "metadata": vis.get("metadata", {}),
                        "data": vis.get("data", {}),
                        "csv_data": csv_data,
                        "csv_summary": csv_summary,
                        "csv_html": csv_to_html_table(csv_data) if csv_data else ""
                    }
                    image_infos.append(image_info)
                    logger.info(f"Image copiée depuis {vis['path']} vers {img_path}")
                    if detailed_interpretation:
                        logger.info(f"  -> Interprétation détaillée de {len(detailed_interpretation)} caractères transférée")
                except Exception as e:
                    logger.error(f"Erreur lors de la copie de l'image {vis['path']}: {e}")

    return image_infos

# ======================================================
# Fonctions pour la mise en forme des résultats de régression
# ======================================================

def format_regression_results(regression_text):
    """
    Formatte les résultats de régression pour une meilleure présentation.
    
    Args:
        regression_text: Texte brut des résultats de régression
        
    Returns:
        str: HTML formatté des résultats
    """
    # Vérifier si c'est un résultat de régression OLS
    if "OLS Regression Results" not in regression_text:
        return f"<pre class='results-block'>{regression_text}</pre>"
    
    # Extraire les sections principales
    sections = {}
    
    # Extraire R-squared et Adj. R-squared
    r_squared_match = re.search(r"R-squared:\s+([\d\.]+)", regression_text)
    adj_r_squared_match = re.search(r"Adj. R-squared:\s+([\d\.]+)", regression_text)
    
    sections['r_squared'] = r_squared_match.group(1) if r_squared_match else "N/A"
    sections['adj_r_squared'] = adj_r_squared_match.group(1) if adj_r_squared_match else "N/A"
    
    # Extraire les coefficients et statistiques
    coef_section = re.search(r"==+\s*\n\s*(coef.*?)(?:\n\s*==+|\Z)", regression_text, re.DOTALL)
    
    # Construire un tableau HTML formatté
    html = '<div class="regression-formatted">'
    html += f'<h4>Résultats de la Régression</h4>'
    
    # Ajouter les métriques clés
    html += '<div class="regression-metrics">'
    html += f'<span class="metric"><strong>R²:</strong> {sections["r_squared"]}</span>'
    html += f'<span class="metric"><strong>R² ajusté:</strong> {sections["adj_r_squared"]}</span>'
    html += '</div>'
    
    # Ajouter le tableau des coefficients si disponible
    if coef_section:
        html += '<table class="regression-coefficients">'
        
        # Créer l'en-tête du tableau
        html += '<thead><tr>'
        headers = re.findall(r"(\w+(?:\s+\w+)*)", coef_section.group(1).split('\n')[0])
        for header in headers:
            html += f'<th>{header}</th>'
        html += '</tr></thead>'
        
        # Ajouter les lignes de coefficients
        html += '<tbody>'
        
        lines = coef_section.group(1).split('\n')[1:] # Skip the header line
        for line in lines:
            if not line.strip():
                continue
            
            html += '<tr>'
            cells = re.findall(r"([\w\.\-]+(?:\s+[\w\.\-]+)*)", line)
            for i, cell in enumerate(cells):
                if i == 0:  # Variable name
                    html += f'<td><strong>{cell}</strong></td>'
                else:
                    html += f'<td>{cell}</td>'
            html += '</tr>'
        
        html += '</tbody></table>'
    
    html += '</div>'
    
    return html

# Function to parse regression tables 
def parse_regression_to_html(narrative_text, regression_tables):
    """
    Tente de trouver une table OLS dans le texte et la convertit en HTML <table>.
    Utilise également les tables de régression capturées si disponibles.
    
    Args:
        narrative_text: Texte contenant potentiellement des tables de régression
        regression_tables: Tables de régression capturées
        
    Returns:
        str: HTML des tables de régression ou None si non trouvé/erreur
    """
    # S'il y a des tables de régression capturées, les utiliser prioritairement
    if regression_tables and len(regression_tables) > 0:
        logger.info(f"Utilisation de {len(regression_tables)} tables de régression capturées")
        
        # Créer un tableau HTML pour chaque table de régression
        html_tables = []
        
        for i, table in enumerate(regression_tables):
            title = table.get('title', f'Régression {i+1}')
            r_squared = table.get('metadata', {}).get('r_squared', 'N/A')
            
            html = f"<div class='regression-container'>\n"
            html += f"<h4>{title}</h4>\n"
            
            # Utiliser l'image de la table
            img_path = table.get('path')
            if img_path and os.path.exists(img_path):
                # Utiliser un chemin relatif basé sur la position du fichier HTML
                rel_path = os.path.basename(img_path)
                html += f"<img src='images/{rel_path}' alt='{title}' class='regression-image' />\n"
                logger.info(f"Image de régression ajoutée: images/{rel_path}")
            else:
                logger.warning(f"Chemin d'image non valide pour {title}: {img_path}")
            
            # Ajouter l'interprétation détaillée si disponible
            detailed_interpretation = table.get('detailed_interpretation', '')
            if detailed_interpretation:
                html += f"<div class='regression-interpretation detailed'>\n"
                html += f"<h5>Interprétation économétrique détaillée</h5>\n"
                html += f"{markdown_filter(detailed_interpretation)}\n"
                html += f"</div>\n"
                logger.info(f"Interprétation détaillée de la régression ajoutée pour {title}")
            # Sinon, utiliser l'interprétation standard si disponible
            elif 'interpretation' in table and table['interpretation']:
                html += f"<div class='regression-interpretation'>\n"
                html += f"<h5>Interprétation</h5>\n"
                html += f"{markdown_filter(table['interpretation'])}\n"
                html += f"</div>\n"
            
            # Ajouter les données CSV si disponibles
            csv_data = table.get('csv_data', '')
            if csv_data:
                csv_html = csv_to_html_table(csv_data)
                if csv_html:
                    html += f"<div class='regression-data-table'>\n<h5>Données de la régression</h5>\n{csv_html}\n</div>\n"
            
            # Ajouter les données structurées si disponibles
            if 'data' in table and table['data']:
                coefficients_data = table['data'].get('coefficients', [])
                if coefficients_data:
                    html += "<div class='regression-data'>\n"
                    html += "<h5>Coefficients significatifs</h5>\n"
                    html += "<table class='coefficients-table'>\n"
                    html += "<thead><tr><th>Variable</th><th>Coefficient</th><th>Valeur p</th><th>Significativité</th></tr></thead>\n"
                    html += "<tbody>\n"
                    
                    for coef in coefficients_data:
                        var_name = coef.get('variable', 'N/A')
                        coef_value = coef.get('coef', 'N/A')
                        p_value = coef.get('p_value', 'N/A')
                        
                        # Marquer les coefficients significatifs
                        if p_value != 'N/A':
                            try:
                                p_value_float = float(p_value)
                                is_significant = p_value_float < 0.05
                                tr_class = "significant" if is_significant else ""
                                significance_text = "Significatif" if is_significant else "Non significatif"
                                html += f"<tr class='{tr_class}'>"
                            except ValueError:
                                html += "<tr>"
                                significance_text = "Indéterminé"
                        else:
                            html += "<tr>"
                            significance_text = "Indéterminé"
                        
                        html += f"<td>{var_name}</td><td>{coef_value}</td><td>{p_value}</td><td>{significance_text}</td></tr>\n"
                    
                    html += "</tbody></table>\n"
                    html += "</div>\n"
            
            html += "</div>\n"
            html_tables.append(html)
        
        return "\n".join(html_tables)
    
    # Pour la recherche dans le texte narrative
    match = re.search(r"={10,}\s*\n\s*OLS Regression Results\s*\n={10,}(.*?)\n={10,}", narrative_text, re.DOTALL | re.IGNORECASE)
    
    if match:
        logger.info("Table de régression OLS détectée dans le texte narrative.")
        table_content = match.group(1).strip()
        
        # Utiliser la fonction de formatage améliorée
        return format_regression_results("OLS Regression Results\n" + table_content)
        
    # Fallback : chercher dans le texte narrative
    match = re.search(r"OLS Regression Results\s*\n={10,}(.*?)\n={10,}", narrative_text, re.DOTALL | re.IGNORECASE) # Essayer sans la première ligne de =

    if match:
        logger.info("Table de régression OLS détectée dans le texte narrative.")
        table_content = match.group(1).strip()
        lines = table_content.split('\n')

        html = "<table class='regression-table'>\n"
        in_header_section = True # Les premières lignes sont souvent des clés:valeurs

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line: continue # Ignore les lignes vides

            # Détecter la ligne d'en-tête du tableau principal (ex: coef std err t P>|t| [0.025 0.975])
            # Ceci est une heuristique et pourrait nécessiter un ajustement
            if re.match(r"^\s*coef\s+std err\s+t\s+P>\|t\|", stripped_line, re.IGNORECASE):
                 in_header_section = False
                 # Ajouter la ligne d'en-tête comme <thead>
                 headers = re.split(r'\s{2,}', stripped_line) # Split sur 2+ espaces
                 html += "  <thead>\n    <tr><th></th>" # Colonne pour le nom de la variable
                 for header in headers:
                     html += f"<th>{header.strip()}</th>"
                 html += "</tr>\n  </thead>\n  <tbody>\n"
                 continue # Passer à la ligne suivante
            elif stripped_line.startswith('---'): # Ligne de séparation
                 continue

            if in_header_section:
                 # Section clé:valeur (ex: R-squared: 0.950)
                 parts = stripped_line.split(':', 1)
                 if len(parts) == 2:
                     html += f"  <tr><td style='font-weight:bold;'>{parts[0].strip()}</td><td colspan='99'>{parts[1].strip()}</td></tr>\n" # Colspan arbitraire
                 else: # Ligne qui n'est pas clé:valeur
                     html += f"  <tr><td colspan='99'>{stripped_line}</td></tr>\n"
            else:
                 # Lignes de données du tableau principal
                 # La première colonne est souvent le nom de la variable, le reste sont des chiffres
                 row_data = re.split(r'\s{2,}', stripped_line)
                 html += "    <tr>\n"
                 for i, cell in enumerate(row_data):
                      style = "text-align: right;" if i > 0 else "" # Aligner les chiffres à droite
                      html += f"      <td style='{style}'>{cell.strip()}</td>\n"
                 html += "    </tr>\n"

        if not in_header_section: # S'il y avait un corps de tableau
            html += "  </tbody>\n"
        html += "</table>\n"
        return html
    else:
        logger.info("Aucune table de régression OLS formatée détectée pour conversion HTML.")
        # Fallback: Mettre tout le bloc dans <pre> si on trouve "OLS Regression Results" mais pas le format attendu
        ols_match = re.search(r"(OLS Regression Results[\s\S]*)", narrative_text, re.IGNORECASE)
        if ols_match:
            logger.warning("Table OLS trouvée mais format non reconnu pour <table>. Utilisation de <pre>.")
            return f"<pre class='regression-text'>{ols_match.group(1).strip()}</pre>"
        return None

# ======================================================
# Fonctions de génération de contenu académique
# ======================================================

def generate_enhanced_introduction(agent1_data, agent2_data, synthesis, discussion, conclusion, economic_reasoning, model, backend):
    """
    Génère une introduction académique de qualité article de recherche en intégrant 
    les résultats des trois agents et l'analyse économique.
    
    Args:
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        synthesis: Résumé simple
        discussion: Section discussion
        conclusion: Section conclusion
        economic_reasoning: Raisonnement économique complet
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Introduction académique de type article de recherche
    """
    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    if "current_model" in agent2_data and agent2_data["current_model"] != model:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']}")
        model = agent2_data["current_model"]

    # Récupérer les éléments importants des agents
    user_prompt = agent1_data.get("user_prompt", "")
    original_introduction = agent1_data.get('llm_output', {}).get('introduction', '')
    hypotheses = agent1_data.get('llm_output', {}).get('hypotheses', '')
    literature_review = agent1_data.get('llm_output', {}).get('literature_review', '')
    methodology = agent1_data.get('llm_output', {}).get('methodology', '')
    
    # Récupérer des extraits du travail de l'agent 2
    narrative = agent2_data.get("narrative", "")
    
    # Récupérer des informations sur les visualisations
    visualisations_info = ""
    if "visualisations" in agent2_data:
        vis_count = len(agent2_data["visualisations"])
        reg_count = sum(1 for v in agent2_data["visualisations"] if v.get('type') == 'regression_table')
        chart_count = vis_count - reg_count
        visualisations_info = f"L'analyse comprend {chart_count} visualisations et {reg_count} modèles de régression."
    
    # Créer le prompt pour une introduction académique de qualité article de recherche
    prompt = f"""## RÉDACTION D'UNE INTRODUCTION ACADÉMIQUE DE TYPE ARTICLE DE RECHERCHE

### Problématique de recherche
{user_prompt}

### Éléments conceptuels (Agent 1)
**Introduction originale**: 
{original_introduction[:800]}...

**Revue de littérature**: 
{literature_review[:500]}...

**Hypothèses**: 
{hypotheses[:500]}...

**Méthodologie proposée**: 
{methodology[:500]}...

### Résultats empiriques (Agent 2)
**Narration des résultats**: 
{narrative[:500]}...

**Structure des visualisations**: 
{visualisations_info}

### Analyse synthétique (Agent 3)
**Résumé**: 
{synthesis[:300]}...

**Discussion**: 
{discussion[:300]}...

**Raisonnement économique**: 
{economic_reasoning[:500]}...

**Conclusion**: 
{conclusion[:300]}...

---

**DIRECTIVE**: Rédigez une introduction académique substantielle de type article de recherche économique (800-1000 mots). Cette introduction doit être de très haute qualité académique et remplacera l'introduction originale dans le rapport final.

**STRUCTURE IMPÉRATIVE**:

1. **CONTEXTUALISATION DU SUJET** (2-3 paragraphes)
   - Présentez le contexte global et l'importance du sujet
   - Ancrez la problématique dans les débats contemporains
   - Articulez clairement les enjeux théoriques et pratiques

2. **REVUE DE LITTÉRATURE INTÉGRÉE** (2-3 paragraphes)
   - Synthétisez les travaux fondateurs et récents
   - Identifiez précisément les lacunes dans la littérature existante
   - Positionnez cette recherche par rapport aux connaissances actuelles

3. **PROBLÉMATIQUE ET QUESTIONS DE RECHERCHE** (1 paragraphe)
   - Formulez une question principale claire et précise
   - Articulez 2-3 sous-questions ou dimensions d'analyse

4. **APPROCHE MÉTHODOLOGIQUE ET APERÇU DES RÉSULTATS** (1-2 paragraphes)
   - Présentez succinctement l'approche méthodologique adoptée
   - Mentionnez les principaux résultats sans les détailler intégralement
   - Mettez en avant la contribution spécifique de cette recherche

5. **STRUCTURE DU RAPPORT** (1 court paragraphe)
   - Présentez l'organisation du reste du rapport de façon concise

**EXIGENCES STYLISTIQUES**:
- Style académique formel mais accessible
- Phrases complexes mais claires
- Vocabulaire économique précis
- Ton neutre et objectif
- TEXTE SIMPLE sans mise en forme excessive, titres ou listes à puces

Cette introduction doit être suffisamment substantielle et rigoureuse pour figurer dans un article académique de premier plan en économie.
"""

    try:
        logger.info(f"Génération d'une introduction académique améliorée avec le modèle {model}")
        introduction = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyer les caractères # qui pourraient se trouver dans le texte
        lines = introduction.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Supprimer les symboles # au début des lignes
            cleaned_line = re.sub(r'^#+ *', '', line)
            cleaned_lines.append(cleaned_line)
        
        introduction = '\n'.join(cleaned_lines)
        
        logger.info(f"Introduction académique améliorée générée ({len(introduction)} caractères)")
        return introduction
    except Exception as e:
        logger.error(f"Erreur lors de la génération de l'introduction académique améliorée: {e}")
        # En cas d'erreur, retourner l'introduction originale
        return original_introduction

def generate_abstract_from_introduction(enhanced_introduction, model, backend):
    """
    Génère un abstract (résumé) à partir de l'introduction améliorée.
    
    Args:
        enhanced_introduction: Introduction améliorée
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Abstract généré à partir de l'introduction
    """
    if not enhanced_introduction:
        logger.error("Introduction améliorée non disponible pour générer l'abstract")
        return "Résumé non disponible."
    
    # Créer le prompt pour générer un abstract à partir de l'introduction
    prompt = f"""## GÉNÉRATION D'UN ABSTRACT ACADÉMIQUE

### Introduction complète de l'article
{enhanced_introduction}

---

**DIRECTIVE**: Rédigez un abstract/résumé concis (environ 150-200 mots) qui synthétise parfaitement l'introduction ci-dessus. Cet abstract servira d'entrée en matière au rapport complet.

**CONTENU REQUIS**:
1. Le contexte de l'étude et son importance
2. L'objectif principal et les questions de recherche
3. L'approche méthodologique utilisée
4. Les principaux résultats ou contributions
5. Une brève conclusion ou implication

**EXIGENCES STYLISTIQUES**:
- Style académique concis et direct
- Un seul paragraphe continu
- Phrases claires et précises
- Vocabulaire économique spécialisé mais accessible
- Aucune citation ou référence spécifique
- Texte simple sans mise en forme

Cet abstract doit donner au lecteur une vision complète et précise de la recherche en un coup d'œil.
"""

    try:
        logger.info(f"Génération d'un abstract à partir de l'introduction améliorée avec le modèle {model}")
        abstract = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyer les caractères # qui pourraient se trouver dans le texte
        lines = abstract.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Supprimer les symboles # au début des lignes
            cleaned_line = re.sub(r'^#+ *', '', line)
            cleaned_lines.append(cleaned_line)
        
        abstract = '\n'.join(cleaned_lines)
        
        logger.info(f"Abstract généré à partir de l'introduction ({len(abstract)} caractères)")
        return abstract
    except Exception as e:
        logger.error(f"Erreur lors de la génération de l'abstract: {e}")
        return "Résumé non disponible. Une erreur s'est produite lors de la génération."

def generate_comprehensive_economic_analysis(agent1_data, agent2_data, model, backend):
    """
    Génère un raisonnement économique complet en agrégeant toutes les interprétations
    des visualisations, tables OLS et la problématique initiale.
    
    Args:
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Raisonnement économique complet
    """
    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    if "current_model" in agent2_data and agent2_data["current_model"] != model:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']}")
        model = agent2_data["current_model"]

    # Récupérer la problématique (prompt utilisateur)
    user_prompt = agent1_data.get("user_prompt", "")
    
    # Récupérer les sections académiques de l'agent 1
    introduction = agent1_data.get('llm_output', {}).get('introduction', '')
    hypotheses = agent1_data.get('llm_output', {}).get('hypotheses', '')
    literature_review = agent1_data.get('llm_output', {}).get('literature_review', '')
    methodology = agent1_data.get('llm_output', {}).get('methodology', '')
    limitations = agent1_data.get('llm_output', {}).get('limitations', '')
    
    # Recueillir toutes les interprétations des visualisations
    all_interpretations = []
    if "visualisations" in agent2_data:
        for i, vis in enumerate(agent2_data["visualisations"]):
            if 'interpretation' in vis and vis['interpretation']:
                title = vis.get('title', f'Visualisation {i+1}')
                interp = vis['interpretation']
                all_interpretations.append(f"### {title}\n{interp}")
    
    # Joindre toutes les interprétations
    interpretations_text = "\n\n".join(all_interpretations)
    
    # Récupérer la narration de l'agent 2
    narrative = agent2_data.get("narrative", "")
    
    # Créer le prompt pour le raisonnement économique complet
    prompt = f"""## Raisonnement économique complet et approfondi

### Problématique de recherche
{user_prompt}

### Contextualisation académique
{introduction[:500]}...

### Hypothèses de recherche
{hypotheses}

### Revue de littérature
{literature_review[:500]}...

### Méthodologie appliquée
{methodology[:500]}...

### Résultats techniques
{narrative[:1000]}...

### Interprétations des visualisations et régressions
{interpretations_text}

---

À partir des éléments ci-dessus, produisez un raisonnement économique complet et approfondi qui inclut les sections suivantes:

1. Synthèse globale
   - Résumez l'objectif et le contexte global de l'étude
   - Rappeler les principales hypothèses testées
   - Résumez les observations empiriques clés

2. Analyse économique approfondie
   - Interprétez les résultats dans un cadre économique rigoureux
   - Identifiez les mécanismes économiques sous-jacents
   - Expliquez les relations causales et corrélations observées
   - Reliez explicitement les résultats aux théories économiques pertinentes
   - Discutez les implications économiques des coefficients significatifs

3. Limites et nuances
   - Examinez la validité interne et externe des résultats
   - Discutez des biais potentiels et de leurs impacts sur l'interprétation
   - Suggérez des perspectives alternatives d'interprétation

4. Implications pratiques et théoriques
   - Formulez des recommandations concrètes pour les décideurs ou acteurs économiques
   - Identifiez les contributions théoriques à la littérature académique

Cette analyse doit être rigoureuse, nuancée et suffisamment développée pour capturer la complexité des phénomènes économiques étudiés (généralement entre 1000 et 2000 mots).

IMPORTANT: 
- Utilisez un format de titre normal (pas de majuscules) et sans pourcentages
- Formattez clairement les points 1, 2, 3, et 4 comme des sections avec des titres
- Structurez les sous-points avec des listes à puces ou des sous-titres
- Utilisez une formulation claire et professionnelle
- Réponse en Francais 
"""

    try:
        logger.info(f"Génération du raisonnement économique complet avec le modèle {model}")
        analysis = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyage supplémentaire pour supprimer les symboles # des titres
        lines = analysis.split('\n')
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            # Si la ligne est un titre (commence par # suivi d'espace), supprimer les #
            if stripped.startswith('# ') or stripped.startswith('## ') or stripped.startswith('### '):
                # Compter le nombre de #
                level = 0
                while level < len(stripped) and stripped[level] == '#':
                    level += 1
                
                # Extraire le titre sans les #
                title_text = stripped[level:].strip()
                
                # Ajouter le titre formaté selon son niveau
                if level == 1:
                    cleaned_lines.append(f"Synthèse globale")
                elif level == 2:
                    cleaned_lines.append(f"Analyse économique approfondie")
                elif level == 3:
                    cleaned_lines.append(f"Limites et nuances") 
                elif level == 4:
                    cleaned_lines.append(f"Implications pratiques et théoriques")
                else:
                    cleaned_lines.append(title_text)
            else:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    except Exception as e:
        logger.error(f"Erreur lors de la génération du raisonnement économique: {e}")
        return "Erreur lors de la génération du raisonnement économique complet."

def generate_comprehensive_visual_analysis(visualizations, agent1_data, agent2_data, model, backend):
    """
    Génère une analyse globale de toutes les visualisations.
    
    Args:
        visualizations: Liste des métadonnées des visualisations avec interprétations
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Analyse globale des visualisations
    """
    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    if "current_model" in agent2_data and agent2_data["current_model"] != model:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']}")
        model = agent2_data["current_model"]

    # Séparer les visualisations des tables de régression
    charts = [v for v in visualizations if v.get('type') != 'regression_table']
    regression_tables = [v for v in visualizations if v.get('type') == 'regression_table']
    
    # Extract visualization titles
    vis_descriptions = []
    for i, vis in enumerate(charts):
        title = vis.get('title', f'Figure {i+1}')
        vis_descriptions.append(f"**{title}**")
        
    # Extract regression table information
    regression_descriptions = []
    for i, table in enumerate(regression_tables):
        title = table.get('title', f'Regression Table {i+1}')
        r_squared = table.get('metadata', {}).get('r_squared', 'N/A')
        variables = table.get('metadata', {}).get('variables', [])
        regression_descriptions.append(f"**{title}** (R² = {r_squared}, Variables: {', '.join(variables[:3])}...)")
    
    # Join visualization descriptions
    vis_descriptions_text = '\n'.join(vis_descriptions)
    regression_descriptions_text = '\n'.join(regression_descriptions)
    
    # Récupérer les interprétations générées par Gemini Flash
    interp_descriptions = []
    for i, vis in enumerate(visualizations):
        title = vis.get('title', f'Item {i+1}')
        interp = vis.get('interpretation', 'Interprétation non disponible')
        interp_descriptions.append(f"### {title}\n{interp}")
    
    interp_text = '\n\n'.join(interp_descriptions)
    
    # Extraire les éléments importants des agents 1 et 2
    introduction = agent1_data.get('llm_output', {}).get('introduction', 'Non disponible')
    hypotheses = agent1_data.get('llm_output', {}).get('hypotheses', 'Non disponible')
    methodology = agent1_data.get('llm_output', {}).get('methodology', 'Non disponible')
    narrative = agent2_data.get("narrative", "Non disponible")
    
    # Créer le prompt pour la synthèse
    prompt = f"""## ANALYSE GLOBALE DES VISUALISATIONS ÉCONOMIQUES

### Question de recherche initiale
{agent1_data.get("user_prompt", "Non disponible")}

### Conceptualisation académique (Agent 1)
Introduction: {introduction[:500]}...
Hypothèses: {hypotheses[:500]}...
Méthodologie: {methodology[:300]}...

### Résultats d'analyse (Agent 2)
Narration: {narrative[:500]}...

### Visualisations à interpréter
{vis_descriptions_text}

### Tables de régression
{regression_descriptions_text}

### Interprétations détaillées existantes
{interp_text[:3000]}

---

Générez une analyse globale CONCISE (250-300 mots maximum) qui:

1. Fait la synthèse des tendances clés observées dans les visualisations
2. Relie directement chaque observation aux hypothèses formulées
3. Se concentre uniquement sur les relations économiques les plus significatives
4. Évite toute description technique au profit d'une interprétation économique claire
5. Utilise un langage simple, factuel et direct
6. En francais 

Cette analyse doit être une vue d'ensemble cohérente et condensée des résultats principaux, sans répéter les détails déjà présents dans les interprétations individuelles.

IMPORTANT: Structurez votre réponse en paragraphes clairs et utilisez des transitions logiques entre les idées.
"""
    
    try:
        logger.info(f"Génération de l'analyse visuelle globale avec le modèle {model}")
        content = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyer les caractères # qui pourraient se trouver dans le texte
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            if line.strip().startswith('#'):
                # Nettoyer les symboles # au début des lignes
                cleaned_line = re.sub(r'^#+ *', '', line)
                cleaned_lines.append(cleaned_line)
            else:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    except Exception as e:
        logger.error(f"Erreur lors de la génération de l'analyse visuelle globale: {e}")
        return "Erreur lors de la génération de l'analyse visuelle globale."

def generate_synthesis(agent1_data, agent2_data, model, backend):
    """
    Génère une synthèse de type résumé simple.
    
    Args:
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Résumé simple
    """
    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    if "current_model" in agent2_data and agent2_data["current_model"] != model:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']}")
        model = agent2_data["current_model"]

    # Récupération des données importantes
    narrative = agent2_data.get("narrative", "Aucune narration disponible")
    
    # Récupération des données académiques de l'agent1 si disponibles
    introduction = agent1_data.get('llm_output', {}).get('introduction', '')
    hypotheses = agent1_data.get('llm_output', {}).get('hypotheses', '')
    methodology = agent1_data.get('llm_output', {}).get('methodology', '')
    
    # Informations sur les visualisations pour enrichir la synthèse
    visualisations_info = ""
    if "visualisations" in agent2_data:
        vis_count = len([v for v in agent2_data["visualisations"] if v.get('type') != 'regression_table'])
        reg_count = len([v for v in agent2_data["visualisations"] if v.get('type') == 'regression_table'])
        visualisations_info = f"L'analyse contient {vis_count} visualisations graphiques et {reg_count} modèles de régression."
    
    prompt = f"""## GÉNÉRATION DE RÉSUMÉ SIMPLE

### Narration détaillée de l'analyse
{narrative[:1500]}

### Contextualisation et objectifs de recherche
{introduction[:500]}

### Hypothèses de recherche
{hypotheses[:300]}

### Méthodologie appliquée
{methodology[:300]}

### Informations sur les visualisations
{visualisations_info}

---

Rédigez un résumé très concis des analyses effectuées 150 mots environ. Cette synthèse doit être facile à comprendre pour des non-spécialistes et se concentrer sur:

1. L'objectif principal de l'analyse (1 phrase)
2. Les principales méthodes utilisées, en termes simples (1 phrase)
3. Les résultats les plus importants (2-3 phrases)
4. La conclusion principale (1 phrase)

IMPORTANT:
- Utilisez un langage simple, direct et clair
- Évitez tout jargon technique
- Structurez le texte en un seul paragraphe sans puces ni énumérations
"""

    try:
        logger.info(f"Appel LLM via backend '{backend}' avec modèle '{model}' pour génération du résumé")
        content = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyer les caractères # qui pourraient se trouver dans le texte
        content = content.replace('# ', '').replace('## ', '').replace('### ', '')
        
        return content
    except Exception as e:
        logger.error(f"Erreur lors de la génération du résumé: {e}")
        return f"Erreur lors de la génération du résumé: {e}"

def generate_discussion_section(agent1_data, agent2_data, model, backend):
    """
    Génère une section discussion simple et accessible.
    
    Args:
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Section discussion
    """
    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    if "current_model" in agent2_data and agent2_data["current_model"] != model:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']}")
        model = agent2_data["current_model"]

    narrative = agent2_data.get("narrative", "Aucune narration disponible")
    hypotheses = agent1_data.get('llm_output', {}).get('hypotheses', '')
    limitations = agent1_data.get('llm_output', {}).get('limitations', '')
    literature_review = agent1_data.get('llm_output', {}).get('literature_review', '')
    
    # Extraire des informations sur les résultats de régression si disponibles
    regression_info = ""
    
    if "visualisations" in agent2_data:
        for vis in agent2_data["visualisations"]:
            if vis.get("type") == "regression_table":
                r_squared = vis.get("metadata", {}).get("r_squared", "N/A")
                variables = vis.get("metadata", {}).get("variables", [])
                regression_info += f"Régression avec R² = {r_squared}, variables clés: {', '.join(variables[:5])}.\n"
                
                # Extraire les interprétations générées par Gemini Flash
                interpretation = vis.get("interpretation", "")
                if interpretation:
                    regression_info += f"Interprétation par Gemini Flash: {interpretation[:200]}...\n\n"
    
    # Récupérer quelques interprétations de visualisations pour enrichir la discussion
    vis_interpretations = []
    if "visualisations" in agent2_data:
        for vis in agent2_data["visualisations"]:
            if vis.get("type") != "regression_table":
                interpretation = vis.get("interpretation", "")
                if interpretation:
                    title = vis.get("title", "Figure")
                    vis_interpretations.append(f"### {title}\n{interpretation[:200]}...")
                    if len(vis_interpretations) >= 2:  # Limiter à 2 interprétations
                        break
    
    vis_interpretations_text = "\n\n".join(vis_interpretations)
    
    prompt = f"""## GÉNÉRATION DE SECTION DISCUSSION CONCISE

### Résultats de l'analyse
{narrative[:1000]}

### Interprétations des visualisations
{vis_interpretations_text}

### Informations sur les régressions 
{regression_info}

### Hypothèses de recherche initiales
{hypotheses[:300]}

### Limitations méthodologiques
{limitations[:300]}

---

Rédigez une section "Discussion" concise (250-300 mots maximum) qui interprète les résultats de l'analyse de façon claire et accessible.

STRUCTURE ESSENTIELLE:
1. **Interprétation des résultats principaux** (1 paragraphe)
   - Expliquez ce que signifient les résultats les plus importants
   - Mettez en évidence les relations les plus significatives

2. **Limites et précautions** (1 paragraphe court)
   - Mentionnez 2-3 limites principales de l'étude
   - Expliquez brièvement leur impact sur l'interprétation

3. **Implications pratiques** (1 paragraphe court)
   - Discutez des applications concrètes de ces résultats

EXIGENCES STYLISTIQUES:
- Langage simple et direct, sans jargon technique
- Paragraphes courts (3-4 phrases maximum)
- Ne dépassez pas 300 mots au total
- Ton neutre et factuel
- Utilisez des connecteurs logiques clairs entre les paragraphes pour assurer une progression fluide des idées
"""

    try:
        logger.info(f"Appel LLM via backend '{backend}' avec modèle '{model}' pour génération de la section discussion")
        content = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyer les caractères # qui pourraient se trouver dans le texte
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            if line.strip().startswith('#'):
                # Nettoyer les symboles # au début des lignes
                cleaned_line = re.sub(r'^#+ *', '', line)
                cleaned_lines.append(cleaned_line)
            else:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    except Exception as e:
        logger.error(f"Erreur lors de la génération de la discussion: {e}")
        return "Erreur lors de la génération de la section discussion."

def generate_conclusion_section(agent1_data, agent2_data, model, backend):
    """
    Génère une conclusion simple.
    
    Args:
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Section conclusion
    """
    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    if "current_model" in agent2_data and agent2_data["current_model"] != model:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']}")
        model = agent2_data["current_model"]

    narrative = agent2_data.get("narrative", "Aucune narration disponible")
    introduction = agent1_data.get('llm_output', {}).get('introduction', '')
    
    # Informations sur les visualisations
    vis_summary = ""
    if "visualisations" in agent2_data:
        vis_types = []
        regression_count = 0
        
        for vis in agent2_data["visualisations"]:
            if vis.get("type") == "regression_table":
                regression_count += 1
            else:
                filename = vis.get("filename", "")
                if "correlation" in filename.lower():
                    vis_types.append("matrices de corrélation")
                elif "scatter" in filename.lower() or "relation" in filename.lower():
                    vis_types.append("nuages de points")
                elif "box" in filename.lower():
                    vis_types.append("boîtes à moustaches")
                elif "histogram" in filename.lower() or "distribution" in filename.lower():
                    vis_types.append("histogrammes")
        
        vis_types = list(set(vis_types))  # Éliminer les doublons
        
        if vis_types or regression_count > 0:
            vis_summary = "L'analyse a utilisé "
            if vis_types:
                vis_summary += f"des visualisations ({', '.join(vis_types)})"
            
            if regression_count > 0:
                if vis_types:
                    vis_summary += f" et {regression_count} modèles de régression"
                else:
                    vis_summary += f"{regression_count} modèles de régression"
            
            vis_summary += " pour explorer les données."
    
    prompt = f"""## GÉNÉRATION DE CONCLUSION CONCISE

### Résultats de l'analyse
{narrative[:800]}

### Introduction et objectifs initiaux
{introduction[:300]}

### Résumé des visualisations
{vis_summary}

---

Rédigez une conclusion très concise (150 mots maximum) pour ce rapport d'analyse. Cette conclusion doit synthétiser les principaux résultats et être facilement compréhensible.

STRUCTURE ESSENTIELLE:
1. **Rappel bref de l'objectif** (1 phrase)
2. **Résumé des 2-3 découvertes les plus importantes** (2-3 phrases)
3. **Implications pratiques principales** (1-2 phrases)
4. **Conclusion générale** (1 phrase d'ouverture sur des perspectives futures)

EXIGENCES:
- Concision extrême: 150 mots maximum
- Langage simple et direct
- Ton positif et constructif
- Pas de nouveaux éléments ou analyses
- Un seul paragraphe compact
- Assurez-vous que la conclusion s'inscrit dans la suite logique des autres sections
"""

    try:
        logger.info(f"Appel LLM via backend '{backend}' avec modèle '{model}' pour génération de la conclusion")
        content = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyer les caractères # qui pourraient se trouver dans le texte
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            if line.strip().startswith('#'):
                # Nettoyer les symboles # au début des lignes
                cleaned_line = re.sub(r'^#+ *', '', line)
                cleaned_lines.append(cleaned_line)
            else:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    except Exception as e:
        logger.error(f"Erreur lors de la génération de la conclusion: {e}")
        return "Erreur lors de la génération de la conclusion."

def generate_references(agent1_data, model, backend):
    """
    Génère des références bibliographiques simples.
    
    Args:
        agent1_data: Données de l'agent1
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        
    Returns:
        str: Références bibliographiques
    """
    literature_review = agent1_data.get('llm_output', {}).get('literature_review', '')
    user_prompt = agent1_data.get('user_prompt', '')
    
    prompt = f"""## GÉNÉRATION DE RÉFÉRENCES SIMPLES

### Indications sur la littérature pertinente
{literature_review[:800]}

### Sujet de recherche
{user_prompt}

---

Générez une liste CONCISE de 4-5 références essentielles pour ce sujet d'analyse. Ces références doivent être authentiques et présentées dans un format simple.

EXIGENCES:
1. **Nombre et variété**: 
   - 4-5 références pertinentes et accessibles
   - Privilégiez les sources récentes (après 2015)

2. **Format simple et compact**:
   - Format: Auteur(s), (année). Titre. Source.
   - Pas de formatage complexe
   - Classement par ordre alphabétique d'auteur

IMPORTANT:
- Les références doivent être RÉELLES et VÉRIFIABLES
- Privilégiez les références en français quand c'est possible
- Évitez les références trop spécialisées ou peu connues
- Utilisez des puces (format Markdown) pour présenter chaque référence
"""

    try:
        logger.info(f"Appel LLM via backend '{backend}' pour génération des références bibliographiques")
        content = call_llm(prompt=prompt, model_name=model, backend=backend)
        
        # Nettoyer les caractères # qui pourraient se trouver dans le texte
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            if line.strip().startswith('#'):
                # Nettoyer les symboles # au début des lignes
                cleaned_line = re.sub(r'^#+ *', '', line)
                cleaned_lines.append(cleaned_line)
            else:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    except Exception as e:
        logger.error(f"Erreur lors de la génération des références: {e}")
        return "Erreur lors de la génération des références bibliographiques."

def interpret_visualization_with_gemini(vis, agent1_data, model, backend, prompts_log_path, timeout=60):
    """
    Interprète une visualisation en demandant une interprétation concise contextualisée avec la prompt utilisateur.
    
    Args:
        vis: Métadonnées de la visualisation avec base64 de l'image
        agent1_data: Données de l'agent1
        model: Modèle LLM à utiliser
        backend: Backend pour les appels LLM
        prompts_log_path: Chemin pour sauvegarder les prompts
        
    Returns:
        str: Interprétation de la visualisation
    """
    from llm_utils import call_llm
    
    # Valider que l'image est disponible
    if 'base64' not in vis or not vis['base64']:
        logger.error(f"Pas de données base64 disponibles pour la visualisation {vis.get('id', 'non identifiée')}")
        return "Erreur: Impossible d'analyser cette visualisation (image non disponible)"
    
    # Extraire le titre et le type
    if 'filename' in vis:
        filename = vis.get("filename", "figure.png")
        vis_id = os.path.splitext(filename)[0]
    else:
        vis_id = vis.get("id", "visualisation")
        filename = vis_id + '.png'
    
    # Déterminer le type de visualisation
    vis_type = "Visualisation"
    if 'regression' in vis_id.lower():
        vis_type = "Table de Régression OLS"
    elif 'correlation' in vis_id.lower() or 'corr' in vis_id.lower():
        vis_type = "Matrice de Corrélation"
    elif 'distribution' in vis_id.lower() or 'hist' in vis_id.lower():
        vis_type = "Graphique de Distribution"
    elif 'scatter' in vis_id.lower() or 'relation' in vis_id.lower():
        vis_type = "Nuage de Points"
    elif 'box' in vis_id.lower():
        vis_type = "Boîte à Moustaches"
    
    # Récupérer le titre
    title = vis.get("title", vis_type)
    
    # Extraire la prompt utilisateur
    user_prompt = agent1_data.get("user_prompt", "")
    
    # Extraire des informations supplémentaires pour le contexte
    extra_context = ""
    if 'metadata' in vis:
        if 'r_squared' in vis['metadata']:
            extra_context += f"\nR-squared: {vis['metadata']['r_squared']}"
        if 'variables' in vis['metadata']:
            extra_context += f"\nVariables principales: {', '.join(vis['metadata']['variables'])}"
    
    # Créer le prompt pour une interprétation concise
    prompt = f"""## INTERPRÉTATION CONCISE DE VISUALISATION

### Type
{vis_type}

### Titre
{title}

### Question utilisateur initiale
"{user_prompt}"

### Données visualisées
Variables: {', '.join(agent1_data["metadata"].get("noms_colonnes", [])[:5])}...
{extra_context}

---

Analyse brièvement cette visualisation en 2-3 phrases maximum. Ton interprétation doit:

1. Identifier les tendances ou relations clés visibles
2. Expliquer comment cette visualisation répond à la question de l'utilisateur
3. Éviter toute description technique de la visualisation elle-même
4. Dans le cas d'un resulat de régression, mentionner les variables clés et leur impact

IMPORTANT: Sois extrêmement concis. Concentre-toi uniquement sur ce qui est pertinent pour répondre à la question utilisateur.
EXCEPTION : Si c'est une table de régression, tu peux interpreter plus longuement les résultats, mais reste dans un format simple et direct.
"""

    try:
        logger.info(f"Appel à Gemini pour interprétation concise de visualisation {vis_id}")
        interpretation = call_llm(
            prompt=prompt, 
            model_name=model, 
            backend=backend, 
            image_base64=vis['base64']
        )
        
        logger.info(f"Interprétation générée pour: {vis_id}")
        return interpretation
    except Exception as e:
        logger.error(f"Erreur lors de l'interprétation de l'image pour {vis_id}: {e}")
        return f"Erreur d'interprétation: {e}"

# ======================================================
# Fonctions principales de génération des rapports
# ======================================================

def generate_report_docx(agent1_data, agent2_data, synthesis, discussion, conclusion, references, economic_reasoning, model_name, user_prompt, img_dir, enhanced_introduction=None):
    """
    Génère un rapport Word au format simple avec python-docx.
    
    Args:
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        synthesis: Résumé simple
        discussion: Section discussion
        conclusion: Section conclusion
        references: Références bibliographiques
        economic_reasoning: Raisonnement économique complet
        model_name: Nom du modèle LLM utilisé
        user_prompt: Prompt initial de l'utilisateur
        img_dir: Répertoire contenant les images
        enhanced_introduction: Introduction académique améliorée (optionnel)
        
    Returns:
        str: Chemin du rapport Word généré ou message d'erreur
    """
    if not DOCX_AVAILABLE:
        logger.error("Module python-docx non disponible. Impossible de générer le rapport Word.")
        return None
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_dir = os.path.join("outputs", f"rapport_{timestamp}")
    os.makedirs(report_dir, exist_ok=True)
    
    # Créer un nouveau document Word
    doc = Document()
    
    # Définir le style de base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Fonction pour convertir le Markdown en texte brut
    h = html2text.HTML2Text()
    h.ignore_links = True
    h.ignore_images = True
    h.ignore_emphasis = True
    h.body_width = 0  # No wrapping
    
    def markdown_to_plain_text(markdown_text):
        """Convertit le Markdown en texte brut"""
        if not markdown_text:
            return ""
        # Supprimer les # des titres
        text = re.sub(r'^#+\s+', '', markdown_text, flags=re.MULTILINE)
        # Supprimer les * et _ d'emphase
        text = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text)
        text = re.sub(r'(\*|_)(.*?)\1', r'\2', text)
        # Supprimer les ``` des blocs de code
        text = re.sub(r'```[\w]*\n(.*?)```', r'\1', text, flags=re.DOTALL)
        # Simplifier les listes
        text = re.sub(r'^\s*[-*+]\s+', '• ', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d+\.\s+', '• ', text, flags=re.MULTILINE)
        
        return text
    
    # Page de titre améliorée
    # Ajouter un titre plus élégant
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(f"Analyse Économique")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    
    # Ajouter un sous-titre avec le prompt utilisateur
    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(f"{user_prompt[:80]}{'...' if len(user_prompt) > 80 else ''}")
    subtitle_run.font.size = Pt(16)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)  # Gris foncé
    
    # Ajouter une ligne horizontale décorative
    border_paragraph = doc.add_paragraph()
    border_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_run = border_paragraph.add_run("―――――――――――――――")
    border_run.font.size = Pt(16)
    border_run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
    
    # Ajouter "Rapport d'analyse économique"
    doc_type_paragraph = doc.add_paragraph()
    doc_type_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_type_run = doc_type_paragraph.add_run("Rapport d'analyse économique")
    doc_type_run.font.size = Pt(14)
    doc_type_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # Gris moyen
    
    # Ajouter la date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y')}")
    
    # Ajouter le modèle utilisé
    model_paragraph = doc.add_paragraph()
    model_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    model_run = model_paragraph.add_run(f"Analyse réalisée avec {model_name}")
    model_run.font.size = Pt(9)
    model_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # Gris moyen
    
    doc.add_page_break()
    
    # Résumé
    heading = doc.add_heading("Résumé", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    doc.add_paragraph(markdown_to_plain_text(synthesis))
    
    # Introduction
    heading = doc.add_heading("Introduction", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    doc.add_paragraph(f"Cette analyse s'intéresse à {user_prompt.lower()}.")
    
    # Utiliser l'introduction améliorée si disponible, sinon utiliser celle de l'agent1
    intro_text = enhanced_introduction if enhanced_introduction else agent1_data.get('llm_output', {}).get('introduction', 'Non disponible')
    doc.add_paragraph(markdown_to_plain_text(intro_text))
    
    # Visualisations et Résultats
    doc.add_page_break()
    heading = doc.add_heading("Visualisations et Résultats", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    doc.add_paragraph("Les visualisations ci-dessous illustrent les relations entre les différentes variables étudiées et permettent d'évaluer les hypothèses formulées dans l'introduction.")
    
    # Traiter les visualisations
    visualisations_data = agent2_data.get("visualisations", [])
    standard_visualizations = [v for v in visualisations_data if v.get('type') != 'regression_table']
    regression_tables = [v for v in visualisations_data if v.get('type') == 'regression_table']
    
    # Parcourir les visualisations standard
    for i, vis in enumerate(standard_visualizations):
        if 'filename' in vis:
            title = vis.get('title', f'Figure {i+1}')
            heading = doc.add_heading(title, level=2)
            heading.style.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
            
            # Ajouter l'image
            img_path = os.path.join(img_dir, vis.get('filename'))
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.error(f"Erreur lors de l'ajout de l'image au document Word: {e}")
            
            # Ajouter l'interprétation
            if 'interpretation' in vis and vis['interpretation']:
                interp_para = doc.add_paragraph()
                interp_para.add_run("Interprétation: ").bold = True
                interp_para.add_run(markdown_to_plain_text(vis['interpretation']))

    # Résultats des régressions
    if regression_tables:
        doc.add_page_break()
        heading = doc.add_heading("Résultats des Régressions", level=1)
        heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
        
        # Texte d'introduction mettant l'accent sur l'importance des régressions
        intro_para = doc.add_paragraph()
        intro_para.add_run("Les modèles de régression présentés ci-dessous constituent ").italic = True
        emphasis_run = intro_para.add_run("le cœur de notre analyse économétrique")
        emphasis_run.italic = True
        emphasis_run.bold = True
        emphasis_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
        intro_para.add_run(". Ils permettent d'analyser de manière rigoureuse les relations multivariées entre les facteurs étudiés et de tester formellement les hypothèses formulées dans l'introduction.").italic = True
        
        # Ajouter une ligne horizontale pour marquer l'importance de cette section
        border_paragraph = doc.add_paragraph()
        border_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        border_run = border_paragraph.add_run("_" * 40)
        border_run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
        
        for i, table in enumerate(regression_tables):
            title = table.get('title', f'Régression {i+1}')
            heading = doc.add_heading(title, level=2)
            heading.style.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
            
            # Ajouter l'image de la régression
            filename = table.get('filename')
            if filename is None:
                # Utiliser l'ID comme fallback ou générer un nom par défaut
                filename = f"{table.get('id', f'regression_{i+1}')}.png"
            img_path = os.path.join(img_dir, filename)
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.error(f"Erreur lors de l'ajout de l'image de régression au document Word: {e}")
            
            # Ajouter l'interprétation détaillée si disponible
            if 'detailed_interpretation' in table and table['detailed_interpretation']:
                heading_interp = doc.add_heading("Interprétation économétrique détaillée", level=3)
                heading_interp.style.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
                
                # Ajout d'un style spécial pour la section d'interprétation
                para_interp = doc.add_paragraph()
                para_interp.style = doc.styles['Normal']
                para_interp.paragraph_format.left_indent = Inches(0.2)
                para_interp.paragraph_format.first_line_indent = Inches(-0.2)
                
                # Ajout d'un symbole pour marquer le début de l'interprétation
                para_interp.add_run("🔍 ").bold = True
                
                # Ajout du contenu de l'interprétation
                para_interp.add_run(markdown_to_plain_text(table['detailed_interpretation']))
            
            # Sinon, utiliser l'interprétation standard si disponible
            elif 'interpretation' in table and table['interpretation']:
                interp_para = doc.add_paragraph()
                interp_para.add_run("Interprétation: ").bold = True
                interp_para.add_run(markdown_to_plain_text(table['interpretation']))
            
            # Ajouter un tableau pour les coefficients significatifs si disponible
            if 'data' in table and table['data'] and 'coefficients' in table['data']:
                coefficients = table['data']['coefficients']
                r_squared = table['data'].get('r_squared', 'N/A')
                
                # Ajouter une ligne pour le R-squared
                r_squared_para = doc.add_paragraph()
                r_squared_para.add_run("R-squared: ").bold = True
                r_squared_para.add_run(r_squared)
                
                # Ajouter un tableau pour les coefficients
                doc.add_heading("Coefficients significatifs", level=4)
                coef_table = doc.add_table(rows=1, cols=4)
                coef_table.style = 'Table Grid'
                
                # En-têtes du tableau
                header_cells = coef_table.rows[0].cells
                header_cells[0].text = "Variable"
                header_cells[1].text = "Coefficient"
                header_cells[2].text = "P-value"
                header_cells[3].text = "Significativité"
                
                # Données des coefficients
                for coef in coefficients:
                    row_cells = coef_table.add_row().cells
                    row_cells[0].text = coef.get('variable', 'N/A')
                    row_cells[1].text = coef.get('coef', 'N/A')
                    p_value = coef.get('p_value', 'N/A')
                    row_cells[2].text = p_value
                    
                    # Déterminer si le coefficient est significatif
                    is_significant = False
                    if p_value != 'N/A':
                        try:
                            is_significant = float(p_value) < 0.05
                        except ValueError:
                            pass
                    
                    row_cells[3].text = "Significatif" if is_significant else "Non significatif"
                    
                    # Mettre en évidence les coefficients significatifs
                    if is_significant:
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
            
            # Ajouter un séparateur entre les régressions
            if i < len(regression_tables) - 1:
                doc.add_paragraph()
                separator = doc.add_paragraph()
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator.add_run("* * *")
                doc.add_paragraph()
    
    # Analyse globale
    doc.add_page_break()
    heading = doc.add_heading("Analyse globale", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    comprehensive_visual_analysis = generate_comprehensive_visual_analysis(
        visualisations_data, 
        agent1_data, 
        agent2_data, 
        model_name, 
        "ollama" if "ollama" in model_name.lower() else "gemini"
    )
    doc.add_paragraph(markdown_to_plain_text(comprehensive_visual_analysis))
    
    # Raisonnement économique approfondi
    doc.add_page_break()
    heading = doc.add_heading("Raisonnement économique approfondi", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    
    # Convertir le raisonnement économique en sections formatées
    eco_reasoning_lines = markdown_to_plain_text(economic_reasoning).split('\n')
    current_level = 1
    for line in eco_reasoning_lines:
        line = line.strip()
        if not line:
            continue
            
        # Détecter si c'est un titre
        if line.startswith("Synthèse globale"):
            heading = doc.add_heading("Synthèse globale", level=2)
            heading.style.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
            current_level = 2
        elif line.startswith("Analyse économique approfondie"):
            heading = doc.add_heading("Analyse économique approfondie", level=2)
            heading.style.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
            current_level = 2
        elif line.startswith("Limites et nuances"):
            heading = doc.add_heading("Limites et nuances", level=2)
            heading.style.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
            current_level = 2
        elif line.startswith("Implications pratiques et théoriques"):
            heading = doc.add_heading("Implications pratiques et théoriques", level=2)
            heading.style.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # Bleu clair
            current_level = 2
        # Détecter les listes
        elif line.startswith('• '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
        # Sinon c'est un paragraphe normal
        else:
            doc.add_paragraph(line)
    
    # Discussion
    doc.add_page_break()
    heading = doc.add_heading("Discussion", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    doc.add_paragraph("Cette section interprète les résultats de l'analyse et discute leurs implications plus générales, en s'appuyant sur les figures présentées et les modèles statistiques développés.")
    doc.add_paragraph(markdown_to_plain_text(discussion))
    
    # Conclusion
    heading = doc.add_heading("Conclusion", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    doc.add_paragraph(markdown_to_plain_text(conclusion))
    
    # Références
    doc.add_page_break()
    heading = doc.add_heading("Références", level=1)
    heading.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Bleu
    
    # Traiter les références ligne par ligne
    refs_clean = markdown_to_plain_text(references)
    for line in refs_clean.split('\n'):
        line = line.strip()
        if line.startswith('• '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            doc.add_paragraph(line)
    
    # Pied de page
    footer_paragraph = doc.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.add_run(f"Rapport généré avec {model_name} | {datetime.now().strftime('%d/%m/%Y')}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # Gris moyen
    
    # Enregistrer le document
    docx_path = os.path.join(report_dir, "rapport.docx")
    try:
        doc.save(docx_path)
        logger.info(f"Rapport Word généré avec succès: {docx_path}")
        return docx_path
    except Exception as e:
        logger.error(f"Erreur lors de la génération du document Word: {e}")
        return f"Erreur: Échec de la génération du document Word. Détails: {e}"

def generate_report_pdf(agent1_data, agent2_data, synthesis, discussion, conclusion, references, economic_reasoning, model_name, user_prompt, report_dir, img_dir, enhanced_introduction=None):
    """
    Génère un rapport PDF au format amélioré avec WeasyPrint et Jinja2.
    
    Args:
        agent1_data: Données de l'agent1
        agent2_data: Données de l'agent2
        synthesis: Résumé simple
        discussion: Section discussion
        conclusion: Section conclusion
        references: Références bibliographiques
        economic_reasoning: Raisonnement économique complet
        model_name: Nom du modèle LLM utilisé
        user_prompt: Prompt initial de l'utilisateur
        report_dir: Répertoire où générer le rapport
        img_dir: Répertoire où sauvegarder les images
        enhanced_introduction: Introduction académique améliorée (optionnel)
        
    Returns:
        str: Chemin du rapport PDF généré ou message d'erreur
    """
    logger.info(f"Répertoire du rapport: {report_dir}")

    # --- Préparation des données pour le template ---
    narrative_md = agent2_data.get("narrative", "Aucune narration disponible.")
    
    # Utiliser l'introduction améliorée si disponible, sinon utiliser celle de l'agent1
    introduction_text = enhanced_introduction if enhanced_introduction else agent1_data.get('llm_output', {}).get('introduction', 'Non disponible')
    
    # Traiter les visualisations (prend en charge à la fois les visualisations standard et les tables de régression)
    visualisations_data = agent2_data.get("visualisations", [])
    
    # Log pour le débogage
    logger.info(f"Nombre de visualisations dans agent2_data: {len(visualisations_data)}")
    for i, vis in enumerate(visualisations_data):
        logger.info(f"Visualisation {i+1}: type={vis.get('type', 'non spécifié')}, base64={'présent' if 'base64' in vis else 'absent'}, path={vis.get('path', 'non spécifié')}")
        logger.info(f"  Données structurées: {'présentes' if 'data' in vis and vis['data'] else 'absentes'}")
        logger.info(f"  Données CSV: {'présentes' if 'csv_data' in vis and vis['csv_data'] else 'absentes'}")
        logger.info(f"  Interprétation: {'présente' if 'interpretation' in vis and vis['interpretation'] else 'absente'}")
    
    # Séparer les visualisations et les tables de régression
    standard_visualizations = [v for v in visualisations_data if v.get('type') != 'regression_table']
    regression_tables = [v for v in visualisations_data if v.get('type') == 'regression_table']
    
    # Sauvegarder les images et obtenir leurs infos (filename, title)
    image_infos = save_images(visualisations_data, img_dir)
    
    logger.info(f"Nombre d'images sauvegardées: {len(image_infos)}")
    for i, img in enumerate(image_infos):
        logger.info(f"Image sauvegardée {i+1}: filename={img.get('filename')}, type={img.get('type')}, size={img.get('size', 'inconnu')}")
        logger.info(f"  Données structurées: {'présentes' if 'data' in img and img['data'] else 'absentes'}")
        logger.info(f"  Données CSV: {'présentes' if 'csv_data' in img and img['csv_data'] else 'absentes'}")
        logger.info(f"  Interprétation: {'présente' if 'interpretation' in img and img['interpretation'] else 'absente'}")

    # Séparer les informations sur les tables de régression et les visualisations standard
    standard_vis_infos = [img for img in image_infos if img.get('type') != 'regression_table']
    regression_table_infos = [img for img in image_infos if img.get('type') == 'regression_table']

    # Extraire et formater la table de régression à partir du texte narrative
    regression_table_html = parse_regression_to_html(narrative_md, regression_table_infos)

    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    if "current_model" in agent2_data and agent2_data["current_model"] != model_name:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']} au lieu de {model_name}")
        model_name = agent2_data["current_model"]

    # Generate a comprehensive visual analysis if needed
    comprehensive_visual_analysis = ""
    if len(image_infos) > 1:  # Only generate if we have multiple visualizations
        logger.info("Génération d'une analyse visuelle globale")
        # Déterminer le backend basé sur le modèle
        backend = "ollama" if "ollama" in model_name.lower() else "gemini"
        comprehensive_visual_analysis = generate_comprehensive_visual_analysis(
            image_infos, 
            agent1_data, 
            agent2_data, 
            model_name, 
            backend
        )

    # --- Création du HTML de rapport avec un template amélioré ---
    # Nouveau template HTML avec les améliorations demandées
    html_template = """<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="UTF-8">
  <title>{{ report_title }}</title>
  <style>
    /* Style de base simplifié */
    body {
      font-family: Arial, sans-serif;
      font-size: 10pt;
      line-height: 1.3;
      color: #333;
      margin: 1.5cm;
    }
    
    /* En-tête et pied de page */
    @page {
      size: A4;
      margin: 2cm 1.5cm;
      @top-center {
        content: "{{ report_title | replace('"', '') }}";
        font-family: Arial, sans-serif;
        font-size: 8pt;
        color: #666;
      }
      @bottom-center {
        content: "Page " counter(page) " / " counter(pages);
        font-family: Arial, sans-serif;
        font-size: 8pt;
        color: #666;
      }
    }
    
    /* Typographie simple */
    h1 {
      font-size: 14pt;
      margin-top: 1em;
      margin-bottom: 0.5em;
      color: #000;
      border-bottom: 1px solid #ccc;
      padding-bottom: 3px;
    }
    
    h2 {
      font-size: 12pt;
      margin-top: 0.8em;
      margin-bottom: 0.4em;
      color: #000;
    }
    
    h3 {
      font-size: 11pt;
      margin-top: 0.7em;
      margin-bottom: 0.3em;
      color: #000;
    }
    
    h4 {
      font-size: 10pt;
      margin-top: 0.6em;
      margin-bottom: 0.3em;
      color: #000;
    }
    
    p {
      margin: 0.4em 0 0.8em 0;
      text-align: justify;
    }
    
    /* Images */
    img {
      max-width: 100%;
      height: auto;
      display: block;
      margin: 0.8em auto;
    }
    
    /* Tableaux simplifiés */
    table {
      width: 100%;
      border-collapse: collapse;
      margin: 0.8em 0;
      font-size: 9pt;
    }
    
    table th {
      background-color: #f2f2f2;
      font-weight: bold;
      text-align: left;
      padding: 4px;
      border: 1px solid #ddd;
    }
    
    table td {
      padding: 4px;
      border: 1px solid #ddd;
    }
    
    /* Listes */
    ul, ol {
      margin: 0.5em 0 0.8em 1.2em;
      padding-left: 0;
    }
    
    li {
      margin-bottom: 0.2em;
    }
    
    /* Saut de page */
    .page-break-before {
      page-break-before: always;
    }
    
    /* Pied de page simple */
    .footer {
      text-align: center;
      font-size: 8pt;
      color: #666;
      margin-top: 1em;
      padding-top: 5px;
      border-top: 1px solid #ddd;
    }

    /* Tableaux CSV */
    .csv-data-table {
      font-size: 8pt;
    }
    
    /* Suppression des styles complexes, ne garder que l'essentiel */
    .cover-page {
      margin-bottom: 2cm;
    }
    
    .cover-title-main {
      font-size: 18pt;
      font-weight: bold;
      text-align: center;
      margin-bottom: 0.5cm;
    }
    
    .cover-title-secondary {
      font-size: 14pt;
      text-align: center;
      margin-bottom: 1cm;
    }
    
    .cover-date {
      font-size: 10pt;
      text-align: center;
      margin-top: 1cm;
    }
    
    .cover-model {
      font-size: 9pt;
      text-align: center;
      color: #666;
      margin-top: 2cm;
    }
    
    /* Figure container simplifié */
    .figure-container {
      margin: 1em 0;
    }
    
    .figure-title {
      font-weight: bold;
      font-size: 10pt;
      text-align: center;
      margin: 0.4em 0;
    }
    
    .figure-interpretation {
      margin: 0.4em 0;
      font-style: italic;
      font-size: 9pt;
    }
  </style>
</head>
<body>
  <!-- Page de couverture simplifiée -->
  <div class="cover-page">
    <div class="cover-title-main">Analyse Économique</div>
    <div class="cover-title-secondary">{{ user_prompt }}</div>
    <div class="cover-date">Généré le {{ generation_date }}</div>
    <div class="cover-model">Analyse réalisée avec {{ model_name }}</div>
  </div>

  <!-- Contenu principal -->
  <!-- Résumé -->
  <h1>Résumé</h1>
  {{ synthesis | markdown | safe }}

  <!-- Introduction -->
  <h1>Introduction</h1>
  <p>Cette analyse s'intéresse à {{ user_prompt | lower }}.</p>
  {{ introduction_text | markdown | safe }}

  <!-- Visualisations -->
  <h1 class="page-break-before">Visualisations et Résultats</h1>
  <p>Les visualisations ci-dessous illustrent les relations entre les différentes variables étudiées.</p>
  
  {% for vis in standard_vis_infos %}
  <h2>{{ vis.title }}</h2>
  <div class="figure-container">
    <img src="images/{{ vis.filename }}" alt="{{ vis.title }}">
  </div>
  {% if vis.interpretation %}
  <p class="figure-interpretation">{{ vis.interpretation | markdown | safe }}</p>
  {% endif %}
  {% endfor %}
  
  <!-- Résultats des régressions -->
  {% if regression_table_infos %}
  <h1 class="page-break-before">Résultats des Régressions</h1>
  
  {% for table in regression_table_infos %}
  <h2>{{ table.title }}</h2>
  <div class="figure-container">
    <img src="images/{{ table.filename }}" alt="{{ table.title }}">
  </div>

  {% if table.detailed_interpretation %}
  <h3>Interprétation économétrique détaillée</h3>
  {{ table.detailed_interpretation | markdown | safe }}
  {% elif table.interpretation %}
  <p class="figure-interpretation">{{ table.interpretation | markdown | safe }}</p>
  {% endif %}

  {% if table.data and table.data.coefficients %}
  <h3>Coefficients significatifs</h3>
  <table>
    <thead>
      <tr>
        <th>Variable</th>
        <th>Coefficient</th>
        <th>p-value</th>
        <th>Significativité</th>
      </tr>
    </thead>
    <tbody>
      {% for coef in table.data.coefficients %}
      {% set is_significant = coef.p_value|float < 0.05 if coef.p_value != 'N/A' else False %}
      <tr>
        <td>{{ coef.variable }}</td>
        <td>{{ coef.coef }}</td>
        <td>{{ coef.p_value }}</td>
        <td>{{ "Significatif" if is_significant else "Non significatif" }}</td>
      </tr>
      {% endfor %}
    </tbody>
  </table>
  {% endif %}
  {% endfor %}
  {% endif %}

  <!-- Analyse globale -->
  {% if comprehensive_visual_analysis %}
  <h1 class="page-break-before">Analyse globale</h1>
  {{ comprehensive_visual_analysis | markdown | safe }}
  {% endif %}

  <!-- Raisonnement économique -->
  <h1 class="page-break-before">Raisonnement économique</h1>
  {{ economic_reasoning | markdown | safe }}

  <!-- Discussion -->
  <h1 class="page-break-before">Discussion</h1>
  {{ discussion_text | markdown | safe }}

  <!-- Conclusion -->
  <h1>Conclusion</h1>
  {{ conclusion_text | markdown | safe }}

  <!-- Références -->
  <h1 class="page-break-before">Références</h1>
  {{ references | markdown | safe }}

  <!-- Footer -->
  <div class="footer">
    Rapport généré avec {{ model_name }} | {{ generation_date }}
  </div>
</body>
</html>"""

    # --- Configuration de Jinja2 ---
    # Créer un environnement Jinja2 à partir d'une chaîne
    env = Environment(loader=FileSystemLoader('.'))
    env.filters['basename'] = basename_filter
    env.filters['markdown'] = markdown_filter
    
    # Créer un template à partir de la chaîne HTML
    template = env.from_string(html_template)

    # --- Données à passer au template ---
    template_data = {
        "report_title": f"Analyse de {user_prompt[:50]}{'...' if len(user_prompt) > 50 else ''}",
        "generation_date": datetime.now().strftime("%d/%m/%Y"),
        "metadata": agent1_data.get("metadata", {}),
        "introduction_text": introduction_text,
        "standard_vis_infos": standard_vis_infos,
        "regression_table_infos": regression_table_infos,
        "comprehensive_visual_analysis": comprehensive_visual_analysis,
        "economic_reasoning": economic_reasoning,
        "synthesis": synthesis,
        "discussion_text": discussion,
        "conclusion_text": conclusion,
        "references": references,
        "model_name": model_name,
        "user_prompt": user_prompt
    }

    # --- Génération du HTML final ---
    try:
        final_html = template.render(template_data)
        html_path = os.path.join(report_dir, "rapport.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(final_html)
        logger.info(f"HTML généré: {html_path}")

    except Exception as e:
        logger.error(f"Erreur lors du rendu du template: {e}")
        return f"Erreur: Le rendu du template a échoué. Détails: {e}"

    # --- Génération du PDF avec les styles améliorés ---
    pdf_path = os.path.join(report_dir, "rapport.pdf")
    try:
        # Style CSS pour le PDF
        enhanced_css = CSS(string="""
            @page { 
                size: A4; 
                margin: 2cm 1.5cm;
                @top-center {
                    content: string(heading);
                    font-family: Arial, sans-serif;
                    font-size: 8pt;
                    color: #666;
                }
                @bottom-center {
                    content: "Page " counter(page) " / " counter(pages);
                    font-family: Arial, sans-serif;
                    font-size: 8pt;
                    color: #666;
                }
            }

            h1 { string-set: heading content() }

            /* Optimisations de base */
            body {
                font-family: Arial, sans-serif;
                font-size: 10pt;
                line-height: 1.3;
                margin: 0;
                color: #333;
            }

            p {
                margin-bottom: 0.5em;
                text-align: justify;
            }

            h1, h2, h3, h4 {
                page-break-after: avoid;
                margin-bottom: 0.5em;
            }

            /* Éviter les sauts de page entre titres et contenu */
            h1 + p, h2 + p, h3 + p {
                page-break-before: avoid;
            }

            /* Simplicité pour les figures */
            .figure-container {
                margin: 1em 0;
                page-break-inside: avoid;
            }

            /* Forcer les sauts de page uniquement quand nécessaire */
            .page-break-before {
                page-break-before: always;
            }

            /* Optimisations pour les tableaux */
            table {
                margin-bottom: 0.8em;
                page-break-inside: avoid;
            }

            /* Meilleure gestion des listes */
            ul, ol {
                margin-bottom: 0.6em;
            }""")

        # Options pour WeasyPrint
        html_obj = HTML(string=final_html, base_url=report_dir)
        html_obj.write_pdf(pdf_path, stylesheets=[enhanced_css])
        
        logger.info(f"Rapport PDF généré avec succès: {pdf_path}")
        return pdf_path
    except Exception as e:
        logger.error(f"Erreur lors de la génération du PDF avec WeasyPrint: {e}")
        try:
            # Tentative simplifiée sans le CSS amélioré
            HTML(string=final_html, base_url=report_dir).write_pdf(pdf_path)
            logger.info(f"Rapport PDF généré (mode dégradé): {pdf_path}")
            return pdf_path
        except Exception as e2:
            logger.error(f"Échec de la génération PDF même en mode dégradé: {e2}")
            return f"Erreur: Échec de la génération PDF. Détails: {e} / {e2}"

def main():
    """
    Fonction principale qui orchestre la génération du rapport.
    """
    parser = argparse.ArgumentParser(description="Agent 3: Synthèse et Rapport PDF")
    parser.add_argument("agent1_output", help="Fichier JSON généré par l'agent 1")
    parser.add_argument("agent2_output", help="Fichier JSON généré par l'agent 2")
    parser.add_argument("user_prompt", help="Prompt utilisateur original")
    parser.add_argument("--model", default="gemma3:27b", help="Modèle LLM à utiliser")
    parser.add_argument("--backend", default="ollama", choices=["ollama", "gemini"], help="Backend LLM à utiliser")
    parser.add_argument("--log-file", help="Fichier de log spécifique")  # Nouvel argument
    args = parser.parse_args()

    # Reconfigurer le logging si un fichier de log spécifique est fourni
    if args.log_file:
        # Supprimer les handlers existants
        for handler in logger.handlers[:]:
            logger.removeHandler(handler)
        
        # Ajouter les nouveaux handlers
        file_handler = logging.FileHandler(args.log_file, mode='a')  # mode 'a' pour append
        file_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
        logger.addHandler(file_handler)
        
        stream_handler = logging.StreamHandler(sys.stderr)
        stream_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
        logger.addHandler(stream_handler)
        
        logger.info(f"Logging redirigé vers {args.log_file}")

    # Lecture des données des agents précédents
    try:
        with open(args.agent1_output, "r", encoding="utf-8") as f:
            agent1_data = json.load(f)
    except Exception as e:
        logger.error(f"Erreur lors de la lecture du fichier de l'agent 1: {e}")
        sys.exit(1)

    try:
        with open(args.agent2_output, "r", encoding="utf-8") as f:
            agent2_data = json.load(f)
    except Exception as e:
        logger.error(f"Erreur lors de la lecture du fichier de l'agent 2: {e}")
        sys.exit(1)

    # Vérifier si un modèle puissant a été utilisé par l'agent 2
    model = args.model
    if "current_model" in agent2_data and agent2_data["current_model"] != args.model:
        logger.info(f"Utilisation du modèle puissant transmis par l'agent 2: {agent2_data['current_model']}")
        model = agent2_data["current_model"]

    # NOUVEAU: Génération de l'introduction académique améliorée
    logger.info(f"Génération de l'introduction académique améliorée avec modèle: {model}")
    enhanced_introduction = generate_enhanced_introduction(
        agent1_data, 
        agent2_data, 
        "", # Synthesis sera générée plus tard
        "", # Discussion sera générée plus tard
        "", # Conclusion sera générée plus tard
        "", # Economic reasoning sera généré plus tard
        model, 
        args.backend
    )

    # Générer le contenu
    logger.info(f"Génération du raisonnement économique complet avec modèle: {model}")
    economic_reasoning = generate_comprehensive_economic_analysis(agent1_data, agent2_data, model, args.backend)

    # NOUVEAU: Génération de l'abstract à partir de l'introduction améliorée
    logger.info(f"Génération de l'abstract à partir de l'introduction améliorée avec modèle: {model}")
    synthesis = generate_abstract_from_introduction(enhanced_introduction, model, args.backend)
    
    logger.info(f"Génération de la section discussion avec modèle: {model}")
    discussion = generate_discussion_section(agent1_data, agent2_data, model, args.backend)
    
    logger.info(f"Génération de la conclusion avec modèle: {model}")
    conclusion = generate_conclusion_section(agent1_data, agent2_data, model, args.backend)
    
    logger.info(f"Génération des références avec modèle: {model}")
    references = generate_references(agent1_data, model, args.backend)

    # Création des répertoires pour les rapports
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_dir = os.path.join("outputs", f"rapport_{timestamp}")
    img_dir = os.path.join(report_dir, "images")
    os.makedirs(report_dir, exist_ok=True)
    os.makedirs(img_dir, exist_ok=True)
    logger.info(f"Répertoire du rapport: {report_dir}")

    # Mise à jour des interprétations des visualisations pour qu'elles soient plus concises
    logger.info("Mise à jour des interprétations de visualisations pour plus de concision")
    
    # Mise à jour des interprétations dans les visualisations
    visualisations_data = agent2_data.get("visualisations", [])
    prompts_log_path = "outputs/prompts_interpretations.txt"
    
    # Dossier temporaire pour sauvegarder les images
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_img_dir = os.path.join("outputs", f"temp_images_{timestamp}")
    os.makedirs(temp_img_dir, exist_ok=True)
    
    # Sauvegarder temporairement les images pour pouvoir les interpréter
    updated_visualisations = []
    
    for vis in visualisations_data:
        if 'base64' in vis:
            try:
                # Générer une interprétation plus concise
                interpretation = interpret_visualization_with_gemini(
                    vis,
                    agent1_data,
                    model,
                    args.backend,
                    prompts_log_path
                )
                
                # Mise à jour de l'interprétation
                vis['interpretation'] = interpretation
                updated_visualisations.append(vis)
                logger.info(f"Interprétation mise à jour pour {vis.get('id', vis.get('filename', 'visualisation'))}")
            except Exception as e:
                logger.error(f"Erreur lors de la mise à jour de l'interprétation: {e}")
                updated_visualisations.append(vis)  # Garder la visualisation même en cas d'erreur
        else:
            logger.warning(f"Pas de données base64 pour {vis.get('id', vis.get('filename', 'visualisation'))}. Interprétation impossible.")
            updated_visualisations.append(vis)
    
    # Mettre à jour les visualisations dans agent2_data
    agent2_data["visualisations"] = updated_visualisations

    logger.info("Génération du rapport PDF")
    pdf_path = generate_report_pdf(
        agent1_data, 
        agent2_data, 
        synthesis, 
        discussion,
        conclusion,
        references,
        economic_reasoning,
        model,
        args.user_prompt,
        report_dir,
        img_dir,
        enhanced_introduction  # Nouvelle introduction améliorée
    )

    # Génération du rapport Word si le module est disponible
    docx_path = None
    if DOCX_AVAILABLE:
        logger.info("Génération du rapport Word")
        docx_path = generate_report_docx(
            agent1_data, 
            agent2_data, 
            synthesis, 
            discussion,
            conclusion,
            references,
            economic_reasoning,
            model,
            args.user_prompt,
            img_dir,
            enhanced_introduction  # Nouvelle introduction améliorée
        )
    else:
        logger.warning("Module python-docx non disponible. Aucun rapport Word ne sera généré.")
    
    # Sortie JSON contenant les chemins et la synthèse
    output = {
        "abstract": synthesis,
        "introduction": enhanced_introduction,
        "discussion": discussion,
        "conclusion": conclusion,
        "references": references,
        "economic_reasoning": economic_reasoning,
        "rapport_pdf": pdf_path if isinstance(pdf_path, str) and pdf_path.endswith(".pdf") else None,
        "rapport_docx": docx_path if isinstance(docx_path, str) and docx_path.endswith(".docx") else None,
        "error": pdf_path if not (isinstance(pdf_path, str) and pdf_path.endswith(".pdf")) else None,
        "model_used": model
    }

    print(json.dumps(output, ensure_ascii=False, indent=2))

    if output["error"]:
        logger.error(f"Le pipeline s'est terminé mais la génération du PDF a échoué: {output['error']}")
        sys.exit(1)
    
    if not DOCX_AVAILABLE:
        logger.warning("La génération du document Word est désactivée car python-docx n'est pas installé")
    elif not output["rapport_docx"]:
        logger.warning("La génération du document Word a échoué")

if __name__ == "__main__":
    main()